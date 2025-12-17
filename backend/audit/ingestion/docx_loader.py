from docx import Document
import io

def ingest_docx(file_bytes: bytes) -> dict:
    """
    Ingests DOCX file.
    Page 1 Text: All paragraphs concatenated.
    Page 2 Rows: Content of the FIRST table found.
    """
    doc = Document(io.BytesIO(file_bytes))
    
    # 1. Extract Text (Page 1)
    # We treat the body text as "Page 1" for header parsing
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    # 2. Extract Table (Page 2)
    # detecting the first table for Related Party logic
    rows_data = []
    if doc.tables:
        table = doc.tables[0]
        # Skip header? Same heuristic: check if first row looks like header
        start_row = 0
        if len(table.rows) > 0:
            header_cells = [c.text.lower() for c in table.rows[0].cells]
            if any("name" in h for h in header_cells):
                start_row = 1
        
        for i in range(start_row, len(table.rows)):
            cells = table.rows[i].cells
            # Normalize to 3 columns
            cell_texts = [c.text.strip() for c in cells]
            if len(cell_texts) < 3:
                cell_texts += [""] * (3 - len(cell_texts))
            
            rows_data.append({
                "business_name": cell_texts[0],
                "criteria_code": cell_texts[1],
                "transaction_type": cell_texts[2]
            })

    return {
        "page_1": {"text": full_text},
        "page_2": {"rows": rows_data}
    }
